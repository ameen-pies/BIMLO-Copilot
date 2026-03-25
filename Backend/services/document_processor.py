import os
import re
import base64
import requests
from typing import List, Dict, Optional
from io import BytesIO


class DocumentProcessor:
    """
    Generic document processor — works with any domain, any language.

    Upgrades over v1:
      - pdfplumber replaces PyPDF2: dramatically better table extraction,
        and exposes embedded images so the vision LLM can describe them.
      - Vision LLM image descriptions: for every image found on a PDF page,
        we call the CF Worker (vision model) to produce a plain-text description.
        That description is injected back into the page text as a
        [IMAGE on page N: <description>] block, making diagrams and figures
        fully searchable just like normal text.
      - DOCX images: same treatment — embedded images in Word docs are
        extracted and described.
      - All other logic (chunking, metadata, LLM classification, reference
        extraction) is unchanged.

    Env vars:
      CF_API_KEY   — Cloudflare Workers AI bearer token
      CF_API_URL   — CF Worker base URL
      VISION_MODEL — model ID to use for image description inside the CF Worker
                     (default: "@cf/llava-hf/llava-1.5-7b-hf")
                     Set this to whatever vision model your worker exposes.
      VISION_MAX_DIM — max image dimension before downscaling (default 1024)
                       Lower = faster + cheaper, higher = more detail.
      SKIP_VISION  — set to "1" to disable vision calls entirely (text-only mode)
    """

    # ── tuneable defaults ─────────────────────────────────────────────────
    _DEFAULT_VISION_MODEL = "@cf/llava-hf/llava-1.5-7b-hf"
    _VISION_PROMPT = (
        "You are analysing a page image or embedded figure from a technical document. "
        "Describe what you see in plain English: components, labels, connections, values, "
        "layout, or any text visible in the image. Be specific and concise (2-5 sentences). "
        "If it is a diagram, wiring schematic, rack layout, floor plan, or table, "
        "describe the structure and key elements."
    )

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size    = chunk_size
        self.chunk_overlap = chunk_overlap
        self._api_key      = os.getenv("CF_API_KEY", "")
        self._api_url      = os.getenv("CF_API_URL", "https://bimloapi.medhelaliamin125.workers.dev")
        self._vision_model = os.getenv("VISION_MODEL", self._DEFAULT_VISION_MODEL)
        self._vision_max_dim = int(os.getenv("VISION_MAX_DIM", "1024"))
        self._skip_vision  = os.getenv("SKIP_VISION", "0").strip() == "1"

    # ── Public entry point ────────────────────────────────────────────────

    def process_document(self, file_path: str) -> List[Dict]:
        """
        Process a document and return text chunks with metadata.
        For PDFs: uses pdfplumber (text + tables + image descriptions).
        For DOCX: uses python-docx (text + image descriptions).
        For TXT:  plain text read.
        """
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            text = self._extract_pdf(file_path)
        elif ext in (".docx", ".doc"):
            text = self._extract_docx(file_path)
        elif ext == ".txt":
            text = self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        metadata = self._extract_metadata(text, file_path)
        return self._create_chunks(text, metadata)

    # ── Text extraction ───────────────────────────────────────────────────

    def _extract_pdf(self, file_path: str) -> str:
        """
        Extract text + tables + image descriptions from a PDF using pdfplumber.

        Per page:
          1. pdfplumber.extract_text()  → body text (much cleaner than PyPDF2)
          2. pdfplumber.extract_tables() → each table formatted as pipe-separated rows
          3. pdfplumber images           → PIL Image → vision LLM → description injected

        Falls back to PyPDF2 if pdfplumber is not installed.
        """
        try:
            import pdfplumber
        except ImportError:
            print("⚠️  pdfplumber not installed — falling back to PyPDF2. "
                  "Run: pip install pdfplumber")
            return self._extract_pdf_pypdf2(file_path)

        full_text_parts: List[str] = []

        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):

                    # ── 1. Body text ──────────────────────────────────────
                    body = page.extract_text() or ""
                    if body.strip():
                        full_text_parts.append(body)

                    # ── 2. Tables ─────────────────────────────────────────
                    tables = page.extract_tables() or []
                    for table in tables:
                        if not table:
                            continue
                        table_lines: List[str] = []
                        for row in table:
                            row_cells = [str(cell).strip() if cell else "" for cell in row]
                            table_lines.append(" | ".join(row_cells))
                        table_text = "\n".join(table_lines)
                        if table_text.strip():
                            full_text_parts.append(f"[TABLE on page {page_num}]\n{table_text}")

                    # ── 3. Images → vision description ────────────────────
                    if not self._skip_vision:
                        image_descriptions = self._describe_page_images(page, page_num)
                        full_text_parts.extend(image_descriptions)

        except Exception as e:
            print(f"Error extracting PDF with pdfplumber: {e}")
            return self._extract_pdf_pypdf2(file_path)

        return "\n\n".join(full_text_parts)

    def _extract_pdf_pypdf2(self, file_path: str) -> str:
        """Legacy PyPDF2 fallback — text only, no tables, no images."""
        try:
            import PyPDF2
            text = ""
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += (page.extract_text() or "") + "\n"
            return text
        except Exception as e:
            print(f"Error extracting PDF (PyPDF2 fallback): {e}")
            return ""

    def _describe_page_images(self, page, page_num: int) -> List[str]:
        """
        Extract images from a pdfplumber page and describe each one via the
        vision LLM.  Returns a list of "[IMAGE on page N: <description>]" strings.

        pdfplumber exposes page.images as a list of dicts; the raw image bytes
        can be reconstructed via the parent PDF's backend.
        """
        descriptions: List[str] = []

        if not self._api_key:
            return descriptions

        raw_images = getattr(page, "images", []) or []
        if not raw_images:
            return descriptions

        for img_idx, img_info in enumerate(raw_images):
            try:
                pil_image = self._pdfplumber_image_to_pil(page, img_info)
                if pil_image is None:
                    continue

                b64 = self._pil_to_b64(pil_image)
                if not b64:
                    continue

                desc = self._call_vision_llm(b64)
                if desc:
                    tag = f"[IMAGE on page {page_num}, figure {img_idx + 1}: {desc}]"
                    descriptions.append(tag)
                    print(f"   🖼️  Described image {img_idx + 1} on page {page_num} ({len(desc)} chars)")

            except Exception as e:
                print(f"   ⚠️  Could not describe image {img_idx + 1} on page {page_num}: {e}")
                continue

        return descriptions

    def _pdfplumber_image_to_pil(self, page, img_info: dict):
        """
        Convert a pdfplumber image dict to a PIL Image.

        pdfplumber stores image metadata but not decoded pixels directly.
        We use the page's underlying pdfminer stream to get the raw bytes,
        then hand them to Pillow for decoding.

        Returns None if extraction fails so callers can skip gracefully.
        """
        try:
            from PIL import Image

            # pdfplumber ≥ 0.7 exposes page.pdf which is the pdfminer PDFDocument
            # Each image dict has 'stream' key with the raw bytes in newer versions
            raw_bytes = img_info.get("stream")

            if raw_bytes is None:
                # Older pdfplumber: try the pdfminer xobject directly
                name = img_info.get("name")
                resources = page.page_obj.get("Resources", {})
                xobjects = resources.get("XObject", {})
                xobj = xobjects.get(name)
                if xobj is None:
                    return None
                xobj.resolve()
                raw_bytes = xobj.get_data()

            if not raw_bytes:
                return None

            img = Image.open(BytesIO(raw_bytes)).convert("RGB")

            # Downscale if needed — keeps vision API costs low
            w, h = img.size
            max_d = self._vision_max_dim
            if max(w, h) > max_d:
                scale = max_d / max(w, h)
                img = img.resize((int(w * scale), int(h * scale)))

            return img

        except Exception:
            return None

    def _extract_docx(self, file_path: str) -> str:
        """
        Extract text + inline images from a DOCX file.
        Images are described by the vision LLM and injected as [IMAGE: ...] tags.
        """
        try:
            from docx import Document
            doc = Document(file_path)

            parts: List[str] = []

            # Paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)

            # Table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        parts.append(row_text)

            # Inline images
            if not self._skip_vision and self._api_key:
                img_descriptions = self._describe_docx_images(doc)
                parts.extend(img_descriptions)

            return "\n".join(parts)

        except Exception as e:
            print(f"Error extracting DOCX: {e}")
            return ""

    def _describe_docx_images(self, doc) -> List[str]:
        """Extract and describe all inline images in a DOCX document."""
        descriptions: List[str] = []
        try:
            from PIL import Image
            from docx.oxml.ns import qn

            for rel in doc.part.rels.values():
                if "image" not in rel.reltype:
                    continue
                try:
                    img_part = rel.target_part
                    raw = img_part.blob
                    pil_img = Image.open(BytesIO(raw)).convert("RGB")

                    w, h = pil_img.size
                    max_d = self._vision_max_dim
                    if max(w, h) > max_d:
                        scale = max_d / max(w, h)
                        pil_img = pil_img.resize((int(w * scale), int(h * scale)))

                    b64 = self._pil_to_b64(pil_img)
                    if not b64:
                        continue

                    desc = self._call_vision_llm(b64)
                    if desc:
                        descriptions.append(f"[IMAGE in document: {desc}]")
                        print(f"   🖼️  Described DOCX image ({len(desc)} chars)")

                except Exception as e:
                    print(f"   ⚠️  Could not describe DOCX image: {e}")
                    continue

        except Exception as e:
            print(f"   ⚠️  DOCX image extraction failed: {e}")

        return descriptions

    def _extract_txt(self, file_path: str) -> str:
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                with open(file_path, "r", encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"Error extracting TXT: {e}")
                return ""
        return ""

    # ── Vision LLM helpers ────────────────────────────────────────────────

    def _pil_to_b64(self, img) -> Optional[str]:
        """Convert a PIL Image to a JPEG base64 string."""
        try:
            buf = BytesIO()
            img.save(buf, format="JPEG", quality=82)
            return base64.b64encode(buf.getvalue()).decode("utf-8")
        except Exception as e:
            print(f"   ⚠️  PIL→base64 failed: {e}")
            return None

    def _call_vision_llm(self, image_b64: str) -> str:
        """
        Send a base64 image to the CF Worker and get a plain-text description.

        The worker payload uses the "vision" task so the worker can route to
        a vision-capable model (e.g. LLaVA, Llama-3.2-Vision).

        The image is sent as a data URI in the prompt so the worker doesn't
        need a separate image upload endpoint — it just passes it straight to
        the vision model as a multimodal message.

        Returns empty string on any failure so callers degrade gracefully.
        """
        if not self._api_key:
            return ""

        try:
            payload = {
                "prompt": self._VISION_PROMPT,
                "image":  image_b64,          # CF Worker reads this and builds the vision message
                "model":  self._vision_model,
                "max_tokens": 200,
                "temperature": 0.1,
                "task": "vision",             # tells the worker to use a vision model
            }
            headers = {
                "Authorization": f"Bearer {self._api_key}",
                "Content-Type": "application/json",
            }
            resp = requests.post(self._api_url, headers=headers, json=payload, timeout=30)
            if resp.status_code == 200:
                raw = (resp.json().get("response") or "").strip()
                # Sanitise — collapse whitespace, strip surrounding quotes
                raw = re.sub(r"\s+", " ", raw).strip().strip('"').strip("'")
                return raw[:500]  # cap description length
            else:
                print(f"   ⚠️  Vision LLM returned {resp.status_code}: {resp.text[:100]}")
                return ""
        except Exception as e:
            print(f"   ⚠️  Vision LLM call failed: {e}")
            return ""

    # ── Metadata extraction ───────────────────────────────────────────────

    def _extract_metadata(self, text: str, file_path: str) -> Dict:
        metadata = {
            "filename": os.path.basename(file_path),
            "doc_type": self._classify_doc_type(text, file_path),
        }
        ref = self._extract_reference(text)
        if ref:
            metadata["project_ref"] = ref
        return metadata

    def _classify_doc_type(self, text: str, file_path: str) -> str:
        """
        Ask the LLM to classify the document type from a short excerpt.
        Returns a short snake_case label (e.g. 'specification', 'report',
        'invoice', 'manual', 'contract').  Falls back to 'document' if
        the LLM is unavailable or the call fails.
        """
        if not self._api_key:
            return "document"

        snippet = text[:800].strip()
        filename = os.path.basename(file_path)

        prompt = (
            "Classify this document into a short doc_type label (snake_case, max 3 words).\n"
            "Examples: specification, technical_report, invoice, contract, user_manual, "
            "meeting_notes, datasheet, plan, calculation_note, legal_document\n\n"
            f"Filename: {filename}\n"
            f"Content preview:\n{snippet}\n\n"
            "Reply with ONLY the doc_type label. Nothing else."
        )

        try:
            resp = requests.post(
                self._api_url,
                headers={
                    "Authorization": f"Bearer {self._api_key}",
                    "Content-Type": "application/json",
                },
                json={"prompt": prompt, "max_tokens": 20, "task": "classify"},
                timeout=15,
            )
            if resp.status_code == 200:
                label = (resp.json().get("response") or "").strip().lower()
                label = re.sub(r"[^\w]", "_", label).strip("_")
                return label or "document"
        except Exception as e:
            print(f"⚠️  Doc-type classification failed: {e}")

        return "document"

    def _extract_reference(self, text: str) -> str:
        """
        Extract a project/document reference number from the text.
        Language-agnostic — matches the structure, not the language.
        """
        patterns = [
            r"(?:ref(?:erence)?|réf(?:érence)?|dokument(?:nummer)?|n[°o]\.?)\s*[:\-]?\s*([A-Z0-9][A-Z0-9\-_/\.]{2,})",
            r"#\s*([A-Z0-9][A-Z0-9\-_/\.]{2,})",
            r"\b([A-Z]{2,6}[-_]\d{3,}(?:[-_][A-Z0-9]+)*)\b",
        ]
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                return m.group(1)
        return ""

    # ── Chunking ──────────────────────────────────────────────────────────

    def _create_chunks(self, text: str, metadata: Dict) -> List[Dict]:
        """
        Split text into overlapping chunks, preferring sentence boundaries.

        [IMAGE ...] and [TABLE ...] blocks are treated as regular text so
        image descriptions and table content flow naturally into chunks
        alongside the surrounding body text.
        """
        chunks: List[Dict] = []
        text = re.sub(r"\s+", " ", text).strip()
        if not text:
            return chunks

        start     = 0
        chunk_id  = 0
        while start < len(text):
            end = start + self.chunk_size
            if end < len(text):
                # Prefer breaking at sentence boundary within a small lookahead window
                boundary = text.rfind(".", end, min(end + 100, len(text)))
                if boundary != -1:
                    end = boundary + 1

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append({
                    "text":     chunk_text,
                    "chunk_id": chunk_id,
                    "metadata": metadata.copy(),
                })
                chunk_id += 1

            start = end - self.chunk_overlap if end < len(text) else len(text)

        return chunks